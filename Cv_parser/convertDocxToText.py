from docx import Document

def convertDocxToText(path):
	print (path,"paths")
	document = Document(path)
	s=[para.text for para in document.paragraphs]
	print(s,"string")
	return "\n".join([para.text for para in document.paragraphs])


